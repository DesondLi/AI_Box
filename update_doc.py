#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def update_document():
    # 打开现有文档
    doc = Document('/mnt/d/AI盒子/AI盒子选型.docx')
    
    # 配置样式（如果不存在）
    try:
        h2 = doc.styles['Heading 2 Custom']
    except:
        h2 = doc.styles.add_style('Heading 2 Custom', WD_STYLE_TYPE.PARAGRAPH)
        h2.font.name = '微软雅黑'
        h2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        h2.font.size = Pt(15)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(31, 73, 125)
    
    try:
        h3 = doc.styles['Heading 3 Custom']
    except:
        h3 = doc.styles.add_style('Heading 3 Custom', WD_STYLE_TYPE.PARAGRAPH)
        h3.font.name = '微软雅黑'
        h3._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        h3.font.size = Pt(13)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(68, 114, 196)
    
    # 添加分页符
    doc.add_page_break()
    
    # ========== 新增内容：国产竞品对比 ==========
    doc.add_paragraph('国产竞品对比分析', style='Heading 2 Custom')
    
    doc.add_paragraph(
        '在国产算力盒子中，与NVIDIA Jetson AGX Orin (64GB)参数定位最为接近的两款产品为：'
        '华为Atlas 500 Pro（昇腾310P）和比特大陆算能SE7（BM1690）。以下进行详细对比分析。'
    )
    
    # ========== 华为 Atlas 500 Pro 详细介绍 ==========
    doc.add_paragraph('一、华为 Atlas 500 Pro（昇腾 310P）', style='Heading 3 Custom')
    
    intro_huawei = """华为Atlas 500 Pro是华为面向企业级边缘计算场景推出的旗舰级AI边缘设备，搭载昇腾310P处理器，
是目前国产边缘算力盒中性能最强、生态最完善的产品之一。主要面向政务信创、智慧城市、工业制造等对国产化有明确要求的场景。

核心特点：
• 算力高达200+ TOPS，是国产边缘端算力第一梯队产品
• 32GB LPDDR4X大内存，支持多模型并发运行
• 完善的工业级设计：宽温、PoE供电、多串口
• 完整的昇腾CANN生态支持，与华为云架构无缝衔接
• 企业级技术支持与服务保障

典型应用场景：政务信创项目落地、多路视频智能分析、园区安防系统、工业质检、交通管控。
"""
    doc.add_paragraph(intro_huawei)
    
    # ========== 比特大陆 SE7 详细介绍 ==========
    doc.add_paragraph('二、比特大陆 算能 SE7（BM1690）', style='Heading 3 Custom')
    
    intro_bitmain = """比特大陆算能SE7搭载最新的BM1690芯片，是算能家族中定位最高的边缘计算产品。凭借64GB LPDDR5大内存和
256 GB/s的超高内存带宽，在大模型推理场景下甚至超越了NVIDIA Orin。是追求极致性价比的私有化部署首选。

核心特点：
• 64GB LPDDR5统一内存，与Orin同级别内存容量
• 256 GB/s内存带宽，甚至优于Orin的204 GB/s
• 128 TOPS INT8算力，支持DeepSeek-V3等最新大模型
• 算能Sophon生态持续优化，开源社区活跃度高
• 相比同级别NVIDIA产品具有显著的价格优势

典型应用场景：律所级私有化LLM机房、法律行业大模型中心、企业级多Agent协同平台、本地大模型服务集群。
"""
    doc.add_paragraph(intro_bitmain)
    
    # ========== 三款产品详细对比表格 ==========
    doc.add_paragraph('三、三款产品多维度详细对比', style='Heading 3 Custom')
    
    # 核心硬件参数对比表
    comparison_data = [
        ['对比维度', 'NVIDIA Jetson AGX Orin (64GB)', '华为 Atlas 500 Pro (昇腾310P)', '比特大陆 算能 SE7 (BM1690)'],
        ['AI算力 (INT8)', '272 TOPS', '200+ TOPS', '128 TOPS'],
        ['GPU/NPU架构', '2048-core Ampere GPU + 64 Tensor Cores', '昇腾310P 达芬奇架构', 'BM1690 TPU'],
        ['内存容量', '64GB LPDDR5', '32GB LPDDR4X', '64GB LPDDR5'],
        ['内存带宽', '204 GB/s', '68 GB/s', '256 GB/s'],
        ['最大LLM支持', '70B-72B (AWQ量化)', '7B~72B全系列', '70B (DeepSeek全系)'],
        ['流畅运行LLM', '7B~70B满血', '14B/30B/70B全流畅', '8B-14B/70B量化'],
        ['存储扩展', 'NVMe SSD', 'NVMe SSD', 'NVMe SSD'],
        ['典型功耗', '40-60W', '工业级散热', '风扇/高风噪'],
        ['参考价格', '13,000-15,800元', '8,000-15,000元', '22,000元']
    ]
    
    table = doc.add_table(rows=len(comparison_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light List Accent 1'
    
    for i, row_data in enumerate(comparison_data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if i == 0:
                        run.font.bold = True
                    run.font.size = Pt(9.5)
    
    doc.add_paragraph()
    
    # 生态与场景对比表
    doc.add_paragraph('生态与场景对比：', style='Heading 3 Custom')
    
    eco_data = [
        ['对比维度', 'NVIDIA Jetson AGX Orin', '华为 Atlas 500 Pro', '比特大陆 算能 SE7'],
        ['软件生态', 'CUDA/TensorRT生态无敌，几乎所有AI框架原生支持', '昇腾CANN生态，华为系产品深度整合', '算能Sophon生态，开源社区活跃度高'],
        ['框架支持', 'PyTorch, TensorFlow, TensorRT-LLM, vLLM, Ollama等全覆盖', 'MindSpore, CANN, PyTorch适配', '算能SDK, LLaMA.cpp, DeepSeek深度优化'],
        ['国产化程度', '美国产品，信创受限', '完全国产，信创一级目录', '国产芯片，信创支持'],
        ['典型行业', '机器人、具身智能、科研、高端制造', '政务、安防、金融、能源、交通', '法律、金融、企业私有化部署'],
        ['技术支持', 'NVIDIA官方 + 全球开发者社区', '华为企业级服务 + 昇腾社区', '算能官方 + 第三方开发者'],
        ['学习曲线', '极低，资料极其丰富', '中等，需学习昇腾生态', '中等，需熟悉算能工具链'],
        ['硬件接口', '双万兆、PCIe 4.0、CAN、USB3.2、MIPI', '双万兆、多网口、PoE、宽温、多串口', '双万兆、PCIe、USB3.0'],
        ['Hermes适配', 'S级，几乎零成本', 'A级，需昇腾环境适配', 'B级，需算能驱动调优']
    ]
    
    table2 = doc.add_table(rows=len(eco_data), cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Light Grid Accent 2'
    
    for i, row_data in enumerate(eco_data):
        for j, cell_data in enumerate(row_data):
            cell = table2.rows[i].cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if i == 0:
                        run.font.bold = True
                    run.font.size = Pt(9.5)
    
    doc.add_paragraph()
    
    # ========== 选型建议总结 ==========
    doc.add_paragraph('四、选型建议总结', style='Heading 3 Custom')
    
    conclusion = """• 追求极致性能与生态：选择NVIDIA Jetson AGX Orin，CUDA生态无敌，适合研发和机器人场景
• 信创/政务硬性要求：选择华为Atlas 500 Pro，国产化最彻底，企业级服务保障
• 大模型推理优先：比特大陆算能SE7，64GB内存+256GB/s带宽，运行70B大模型体验最佳
• 预算敏感场景：华为Atlas 500 Pro在8000-10000元价位段极具性价比，适合大规模部署
"""
    doc.add_paragraph(conclusion)
    
    doc.add_paragraph(
        '综合来看，华为Atlas 500 Pro是目前国产算力盒子中综合能力最接近Orin的产品，'
        '而算能SE7在大模型推理的内存指标上甚至超越了Orin。三者在不同场景下各有优势。'
    )
    
    # 保存文档
    doc.save('/mnt/d/AI盒子/AI盒子选型_更新版.docx')
    print("文档已更新并保存为: AI盒子选型_更新版.docx")

if __name__ == '__main__':
    update_document()
