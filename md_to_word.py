#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_styled_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.5
    doc.styles['Normal'].paragraph_format.space_after = Pt(6)
    
    # 标题1样式
    h1 = doc.styles.add_style('Heading 1 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h1.font.name = '微软雅黑'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    
    # 标题2样式
    h2 = doc.styles.add_style('Heading 2 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name = '微软雅黑'
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(31, 73, 125)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    
    # 标题3样式
    h3 = doc.styles.add_style('Heading 3 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h3.font.name = '微软雅黑'
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(68, 114, 196)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    # 标题4样式
    h4 = doc.styles.add_style('Heading 4 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h4.font.name = '微软雅黑'
    h4._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(91, 155, 213)
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)
    
    # 代码块样式
    code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_style.font.color.rgb = RGBColor(0, 0, 0)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.line_spacing = 1.2
    
    # 列表项样式
    list_style = doc.styles.add_style('List Item', WD_STYLE_TYPE.PARAGRAPH)
    list_style.font.name = '微软雅黑'
    list_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.space_after = Pt(3)
    
    return doc

def parse_table(lines):
    """解析markdown表格"""
    rows = []
    for line in lines:
        if '|---' in line or line.strip() == '|':
            continue
        # 分割单元格
        cells = [c.strip() for c in line.split('|')]
        # 移除首尾空字符串
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows

def add_table(doc, data):
    """添加美观的表格"""
    if not data:
        return
    
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格样式
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data
            
            # 设置字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if i == 0:  # 表头
                        run.font.bold = True
                        run.font.size = Pt(11)
                    else:
                        run.font.size = Pt(10)
    
    doc.add_paragraph()

def md_to_word(md_file, docx_file):
    doc = create_styled_document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 检测代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                for code_line in code_lines:
                    p = doc.add_paragraph(code_line, style='Code Block')
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 检测表格
        if line.startswith('|'):
            table_lines.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            # 结束表格
            table_data = parse_table(table_lines)
            add_table(doc, table_data)
            table_lines = []
            in_table = False
        
        # 分隔线
        if line.strip() == '---':
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # 标题
        if line.startswith('#### '):
            doc.add_paragraph(line[5:], style='Heading 4 Custom')
            i += 1
            continue
        if line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading 3 Custom')
            i += 1
            continue
        if line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 2 Custom')
            i += 1
            continue
        if line.startswith('# '):
            doc.add_paragraph(line[2:], style='Heading 1 Custom')
            i += 1
            continue
        
        # 列表项
        if line.startswith('- '):
            doc.add_paragraph('• ' + line[2:], style='List Item')
            i += 1
            continue
        
        # 编号列表
        if re.match(r'^\d+\. ', line):
            doc.add_paragraph(line, style='List Item')
            i += 1
            continue
        
        # 空行
        if line.strip() == '':
            i += 1
            continue
        
        # 普通段落
        doc.add_paragraph(line)
        i += 1
    
    # 处理文件末尾可能剩余的表格
    if in_table and table_lines:
        table_data = parse_table(table_lines)
        add_table(doc, table_data)
    
    doc.save(docx_file)
    print(f"文档已保存到: {docx_file}")

if __name__ == '__main__':
    md_to_word('D:/AI盒子/AI盒子选型.md', 'D:/AI盒子/AI盒子选型.docx')
